from docx import Document
from enums import SectionType
from section_data import ResumeDocument, ResumeSection

class ResumeParser:
    """
    A class to parse and structure a resume from a .docx file.
    
    This class will extract all relevant sections of a resume like Education, Experience,
    Skills, Leadership, Activities, and Projects.
    """
    
    def __init__(self, file_path):
        """
        Initializes the ResumeParser with the path to the .docx file.
        
        Args:
            file_path (str): The path to the .docx file containing the resume.
        """
        self.file_path = file_path
        
        # open the docx file using the docx library
        
        self.doc = Document(file_path)  
        
        # Initialize a ResumeDocument to store all parsed data
        
        self.resume_document = ResumeDocument(file_path=file_path)
        
        # Initialize a ResumeSection for each section type in the enum
        
        for section_type in SectionType:
            self.resume_document.add_section(ResumeSection(section_type=section_type.value))
        
    def parse(self) -> dict:
        """
        Parse the resume sections from the provided .docx file.

        This method processes the content of the .docx file and organizes it into different
        sections: Education, Experience, Skills & Interests, Leadership & Activities, and Projects.
        It assumes that these sections are clearly marked in the document.
        """
        
        # Iterate through each paragraph in the document
        
        # A variable to track which section we're currently parsing
        
        current_section = None  
        
        for section in self.doc.paragraphs:
            """  
            It parses the resume sections from the provided .docx file line by line. For example:
            Sachin Kharel
            sachink1630@gmail.com • linkedin.com/in/sachinkharel • github.com/sachinkharel • sachinkharel.com.np
            Education
            University of Tasmania | Masters of Information Technology and Systems            Hobart, Australia | Feb 2024 - Present
            GPA: 5.83/7
            """
            
            # Clean up the text by stripping unnecessary spaces or special characters
            
            text = section.text.strip()
            
            # Check if the paragraph text corresponds to a section header
            
            if text == SectionType.EDUCATION.value:
                current_section = SectionType.EDUCATION
            elif text == SectionType.EXPERIENCE.value:
                current_section = SectionType.EXPERIENCE
            elif text == SectionType.SKILLS_AND_INTERESTS.value:
                current_section = SectionType.SKILLS_AND_INTERESTS
               
            elif text == SectionType.LEADERSHIP_AND_ACTIVITIES.value:
                current_section = SectionType.LEADERSHIP_AND_ACTIVITIES
            elif text == SectionType.PROJECTS.value:
                current_section = SectionType.PROJECTS

                
            
            # If we are in a section and the text is not empty, process the data
            
            elif current_section and text:
                if current_section == SectionType.EDUCATION:
                    
                    # Process education data
                    
                    data = self.parse_education_data(section)
                    if data:
                        self.resume_document.get_section(SectionType.EDUCATION.value).add_entry(data)
                elif current_section == SectionType.EXPERIENCE:
                    
                    # Process experience data
                    
                    data = self.parse_experience_data(section)
                    if data:
                       self.resume_document.get_section(SectionType.EXPERIENCE.value).add_entry(data)
                elif current_section == SectionType.SKILLS_AND_INTERESTS:
                    
                    # Process skills data
                    
                    data = self.parse_skills_data(section)
                    if data:
                        self.resume_document.get_section(SectionType.SKILLS_AND_INTERESTS.value).add_entry(data)
                elif current_section == SectionType.LEADERSHIP_AND_ACTIVITIES:
                    
                    # Process leadership data
                    
                    data = self.parse_leadership_data(section)
                    if data:
                        self.resume_document.get_section(SectionType.LEADERSHIP_AND_ACTIVITIES.value).add_entry(data)
                elif current_section == SectionType.PROJECTS:
                    
                    # Process project data
                    
                    data = self.parse_projects_data(section)
                    if data:
                         self.resume_document.get_section(SectionType.PROJECTS.value).add_entry(data)
        
        # Return the structured data as a dictionary
        
        return self.resume_document.to_dict()  
    
    def extract_hyperlink(self, paragraph):
        """
        Extracts the first hyperlink (if any) from the given paragraph using XML relationships.
        """
        
        for rel in paragraph.part.rels.values():
            if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink':
                if rel.rId in paragraph._p.xml:
                    return rel.target_ref 
        return ""


    
    def parse_education_data(self, paragraph):
        """
        Parses education-related information from the provided text.
        This can be customized as per the document layout.
        """
        # Clean up the text to handle tabs or multiple spaces
        
        text = paragraph.text.replace("\t", " | ")  # Replace any tabs with a pipe to standardize
        
        # print(f"Parsing education data: {text}")  # Debugging output
        
        # Split the text into parts by '|'
        
        parts = [part.strip() for part in text.split(" | ") if part.strip()]    # Clean up empty parts so that we can get the right length and index
        
        # print(f"Split parts: {parts}")  # Debugging output
        
        link = self.extract_hyperlink(paragraph)
        
        # We expect at least 4 parts: institution, degree, location, duration
        
        if len(parts) >= 4:  
            
            # Return the parsed data in a dictionary format
            
            return {
                "institution": parts[0],  
                "degree": parts[1],      
                "location": parts[2],     
                "duration": parts[3],     
                "gpa": parts[4] if len(parts) > 4 else "", 
                "link": link  
            }
        return None

    
    def parse_experience_data(self, paragraph):
        """
        Parses experience-related information from the provided text.
        This can be customized as per the document layout.
        """

        # Format: Company | Role | Location | Duration
        
        # Replace any tabs with a pipe to standardize
        
        text = paragraph.text.replace("\t", " | ")  
        
        # Clean up empty parts so that we can get the right length and index
        
        parts = [part.strip() for part in text.split(" | ") if part.strip()]   

        link = self.extract_hyperlink(paragraph)

        if len(parts) >= 4:
            self.current_experience_entry = {
                "title": parts[0].strip(),
                "company": parts[1].strip(),
                "location": parts[2].strip(),
                "duration": parts[3].strip(),
                "description": [],  
                "link": link  
            }
            return self.current_experience_entry
        
        # If it's a description line and a current entry exists, append the description
        
        if len(parts) == 1 and hasattr(self, 'current_experience_entry'):
            self.current_experience_entry["description"].append(parts[0].strip())
            return None    
        
        return None

    def parse_skills_data(self, paragraph):
        """
        Parses skills & interests-related information from the provided text.
        """
        
        # Assuming skills are listed by category: technical, languages, and interests
        
        categories = paragraph.text.split(" || ")
        data = {}
        
        # Remove headings like "Technical:", "Language:", and "Interests: and split by commas
        
        if "Technical:" in categories[0]:
            data["technical"] = categories[0][10:].split(",")  
        if "Language:" in categories[1]:
            data["languages"] = categories[1][10:].split(",")  
        if "Interests:" in categories[2]:
            data["interests"] = categories[2][10:].split(",")  
        return data

    def parse_leadership_data(self, paragraph):
        """
        Parses leadership and activities-related information from the provided text.
        """
        
        # Replace any tabs with a pipe to standardize
         
        text = paragraph.text.replace("\t", " | ") 
        
        # Clean up empty parts so that we can get the right length and index
        
        parts = [part.strip() for part in text.split(" | ") if part.strip()]    

        link = self.extract_hyperlink(paragraph)
        
        # Check if the text contains a title and other structured data
        
        if len(parts) >= 4:
            
            # Create a new entry for the title and other details
            
            self.current_leadership_entry = {
                "title": parts[0].strip(),
                "organization": parts[1].strip(),
                "location": parts[2].strip(),
                "duration": parts[3].strip(),
                "description": [], 
                "link": link 
            }
            return self.current_leadership_entry
        
        # If it's a description line and a current entry exists, append the description
        
        if len(parts) == 1 and hasattr(self, 'current_leadership_entry'):
            self.current_leadership_entry["description"].append(parts[0].strip())
            return None
        
        return None

    def parse_projects_data(self, paragraph):
        """
        Parses project-related information from the provided text.
        """
        # Replace any tabs with a pipe to standardize
        
        text = paragraph.text.replace("\t", " | ")  
        
        # Format: Title | Duration
        
        parts = [part.strip() for part in text.split(" | ") if part.strip()]    # Clean up empty parts so that we can get the right length and index
        
        link = self.extract_hyperlink(paragraph)
        
        if len(parts) >= 2:
            
            # Initialize a new project entry
            
            self.current_project_entry = {
                "title": parts[0].strip(),
                "duration": parts[1].strip(),
                "description": [], 
                "link": link  
            }
            return self.current_project_entry
        
        # If it's a description line and a current entry exists, append the description
        
        if len(parts) == 1 and hasattr(self, 'current_project_entry'):
            self.current_project_entry["description"].append(parts[0].strip())
            return None
        
        return None
